"""Build a professional DOCX CRS Blueprint from validated structured data.

The builder consumes the canonical structured blueprint produced by
src.structured_blueprint. It does not parse model-generated markdown for layout.
Legacy markdown responses are converted to the structured model before rendering.
"""
from __future__ import annotations

from datetime import datetime, UTC
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .structured_blueprint import (
    SECTION_ORDER,
    SECTION_TITLES,
    normalise_structured_result,
    normalise_evidence_status,
    display_evidence_status,
    run_quality_gate,
    safe_text,
)

BRAND_BLUE = RGBColor(0x0F, 0x2A, 0x5F)
ACCENT = RGBColor(0x4F, 0x46, 0xE5)
MUTED = RGBColor(0x5B, 0x68, 0x7A)
WARNING = RGBColor(0x92, 0x4A, 0x00)
LIGHT_FILL = "F3F6FA"
HEADER_FILL = "E9EEF8"
STATUS_FILL = "FFF7E6"

REVIEW_NOTE = "Verify jurisdiction-specific requirements against official guidance before implementation."
LINKEDIN_URL = "https://www.linkedin.com/in/abhinit-sen-63443015/"

DOC_SECTION_ORDER = SECTION_ORDER + ["fatca", "evidence"]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: Any, *, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(safe_text(text))
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _format_table(table, header_fill: str = HEADER_FILL) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            if row_idx == 0:
                _set_cell_shading(cell, header_fill)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = BRAND_BLUE


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(safe_text(text), level=level)
    if not p.runs:
        p.add_run(safe_text(text))
    for run in p.runs:
        run.font.color.rgb = BRAND_BLUE if level <= 1 else ACCENT
        run.font.bold = True
    p.paragraph_format.space_before = Pt(8 if level > 1 else 12)
    p.paragraph_format.space_after = Pt(4)


def _add_small_note(doc: Document, text: str, color: RGBColor = MUTED) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(safe_text(text))
    run.font.size = Pt(8)
    run.font.color.rgb = color


def _add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable external hyperlink to a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "4F46E5")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = safe_text(text)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_next_step(doc: Document) -> None:
    _add_heading(doc, "Next step", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Need this mapped to your actual systems? Connect with " )
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    _add_hyperlink(p, "Abhinit Sen for a Blueprint Review", LINKEDIN_URL)
    r2 = p.add_run(".")
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED


def _add_status_badge(paragraph, status: str) -> None:
    status = display_evidence_status(status)
    run = paragraph.add_run(f" [{status}]")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = WARNING if status == "Needs verification" else ACCENT


def _add_numbered_list(doc: Document, items: list[dict[str, Any]]) -> None:
    for idx, item in enumerate(items, start=1):
        # Use Word's List Number style and explicit restart-friendly text. This avoids
        # cross-section continuation in headless Word/LibreOffice rendering.
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(safe_text(item.get("text")))
        _add_status_badge(p, item.get("evidence_status"))


def _add_bullet_list(doc: Document, items: list[dict[str, Any]]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(safe_text(item.get("text")))
        _add_status_badge(p, item.get("evidence_status"))


def _add_table_from_rows(doc: Document, columns: list[str], rows: list[dict[str, Any]], *, title: str | None = None) -> None:
    columns = [safe_text(c) for c in columns if safe_text(c)]
    if not columns or not rows:
        return
    if title:
        _add_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=len(columns))
    table.allow_autofit = True
    for c_idx, column in enumerate(columns):
        _set_cell_text(table.rows[0].cells[c_idx], column, bold=True, size=8)
    _set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for c_idx, column in enumerate(columns):
            value = row.get(column, "")
            if column == "Evidence Status":
                value = display_evidence_status(value)
            _set_cell_text(cells[c_idx], value, size=8)
    _format_table(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_block(doc: Document, block: dict[str, Any]) -> None:
    title = block.get("title")
    btype = block.get("type")
    if title and btype not in {"table", "review_table"}:
        _add_heading(doc, title, 3)
    if btype == "numbered_list":
        _add_numbered_list(doc, block.get("items", []))
    elif btype == "bullet_list":
        _add_bullet_list(doc, block.get("items", []))
    elif btype in {"table", "review_table"}:
        _add_table_from_rows(doc, block.get("columns", []), block.get("rows", []), title=title)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(safe_text(block.get("text")))
        run.font.size = Pt(9)
        _add_status_badge(p, block.get("evidence_status"))


def _add_metadata_table(doc: Document, params: dict, quality: dict) -> None:
    rows = [
        ("Document status", "Implementation draft"),
        ("Jurisdiction", params.get("jurisdiction", "")),
        ("Institution type", params.get("fi_type", "")),
        ("Reporting year", params.get("reporting_year", "")),
        ("Generated", datetime.now(UTC).strftime("%d %B %Y %H:%M UTC")),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Metadata", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Value", bold=True)
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, HEADER_FILL)
    for k, v in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], k, bold=True)
        _set_cell_text(cells[1], v)
    _format_table(table)


def _add_quality_table(doc: Document, quality: dict) -> None:
    checks = quality.get("checks", [])
    if not checks:
        return
    _add_heading(doc, "Pre-download Quality Gate", 2)
    table = doc.add_table(rows=1, cols=3)
    headers = ["Check", "Result", "Detail"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for check in checks:
        cells = table.add_row().cells
        _set_cell_text(cells[0], check.get("name", ""))
        _set_cell_text(cells[1], "Pass" if check.get("passed") else "Fail")
        _set_cell_text(cells[2], check.get("detail", ""))
    _format_table(table)


def _add_static_toc(doc: Document, structured: dict) -> None:
    _add_heading(doc, "Table of Contents", 1)
    _add_small_note(doc, "Use Word's Update Field command if an automatic page-numbered TOC is required.")
    for idx, key in enumerate([k for k in DOC_SECTION_ORDER if k in structured.get("sections", {})], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{idx}. {structured['sections'][key].get('title', SECTION_TITLES.get(key, key))}")


def _add_change_history_and_approval(doc: Document) -> None:
    _add_heading(doc, "Change History", 1)
    rows = [
        {"Version": "0.1", "Date": datetime.now(UTC).strftime("%d %B %Y"), "Author": "CRS Blueprint", "Change": "Initial generated draft for review."},
        {"Version": "", "Date": "", "Author": "", "Change": ""},
    ]
    _add_table_from_rows(doc, ["Version", "Date", "Author", "Change"], rows)

    _add_heading(doc, "Approval Block", 1)
    rows = [
        {"Role": "Compliance Owner", "Name": "", "Decision": "Approve / Reject / Return for changes", "Date": ""},
        {"Role": "Legal / Tax Reviewer", "Name": "", "Decision": "Approve / Reject / Return for changes", "Date": ""},
        {"Role": "Technology Owner", "Name": "", "Decision": "Approve / Reject / Return for changes", "Date": ""},
    ]
    _add_table_from_rows(doc, ["Role", "Name", "Decision", "Date"], rows)


def _setup_document_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = BRAND_BLUE if style_name == "Heading 1" else ACCENT
        style.font.bold = True


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def _setup_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("CRS Blueprint | Generated " + datetime.now(UTC).strftime("%d %b %Y"))
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def build_docx(result: dict, params: dict) -> str:
    structured = result.get("_structured_blueprint") if isinstance(result, dict) else None
    if not structured:
        structured = normalise_structured_result(result or {}, params, {})
    quality = result.get("_quality_gate") if isinstance(result, dict) else None
    if not quality:
        quality = run_quality_gate(structured)

    doc = Document()
    _setup_document_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading("CRS Blueprint", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BRAND_BLUE
        run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Implementation requirements draft")
    r.bold = True
    r.font.size = Pt(11)

    _add_metadata_table(doc, params, quality)
    disc_p = doc.add_paragraph()
    disc_p.paragraph_format.space_before = Pt(4)
    disc_p.paragraph_format.space_after = Pt(2)
    disc_run = disc_p.add_run("Review note: " + REVIEW_NOTE)
    disc_run.font.size = Pt(8)
    disc_run.font.color.rgb = WARNING

    _add_static_toc(doc, structured)

    _add_change_history_and_approval(doc)

    # Implementation sections are table-heavy; landscape pages make the
    # field catalogue, controls, tests and verification tasks readable.
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_landscape(body_section)

    for idx, key in enumerate([k for k in DOC_SECTION_ORDER if k in structured.get("sections", {})], start=1):
        section = structured["sections"][key]
        _add_heading(doc, f"{idx}. {section.get('title', SECTION_TITLES.get(key, key))}", 1)
        for block in section.get("blocks", []):
            _add_block(doc, block)
        # Sections intentionally flow together to avoid sparse, footer-only pages.

    if "evidence" not in structured.get("sections", {}):
        _add_heading(doc, "Evidence, Assumptions and Review", 1)

    _add_next_step(doc)

    _setup_footer(doc)

    jur_slug = safe_text(params.get("jurisdiction", "unknown")).replace(" ", "_") or "unknown"
    filename = f"CRS_Blueprint_{jur_slug}_{datetime.now(UTC).strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename
